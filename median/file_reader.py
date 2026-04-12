import io
import re

import pypdf
from docx import Document

from median.utils import median_logger


def clean_extracted_text(content: str | None) -> str:
    """Normalizes whitespace so downstream card generation gets cleaner text."""

    if not content:
        return ""

    content = content.replace("\x00", " ")
    content = re.sub(r"\r\n?", "\n", content)
    content = re.sub(r"[ \t]+", " ", content)
    content = re.sub(r"\n{3,}", "\n\n", content)
    return content.strip()


def read_docx(docx_file):
    """
    Reads the content of a DOCX file.

    Args:
        docx_file: The path to the DOCX file or a file-like object.

    Returns:
        str: The content of the DOCX file as a string.
    """

    if isinstance(docx_file, str):
        doc = Document(docx_file)
    else:
        doc = Document(io.BytesIO(docx_file.read()))
    return clean_extracted_text("\n".join(paragraph.text for paragraph in doc.paragraphs))


def read_pdf(pdf_file):
    """
    Reads the content of a PDF file.

    Args:
        pdf_file: The path to the PDF file or a file-like object.

    Returns:
        str: The content of the PDF file as a string, extracted from all pages.
    """

    try:
        if isinstance(pdf_file, str):
            with open(pdf_file, "rb") as f:
                pdf_reader = pypdf.PdfReader(f)
        else:
            pdf_reader = pypdf.PdfReader(io.BytesIO(pdf_file.read()))

        pages_text = [
            pdf_reader.pages[page].extract_text() or ""
            for page in range(len(pdf_reader.pages))
        ]
        return clean_extracted_text("\n\n".join(pages_text))
    except Exception as e:  # Consider catching more specific exceptions
        median_logger.error(f"Error reading PDF file: {e}")
        return ""


def main(file, file_type):
    """
    Main function to read the content of different file types based on the specified file type.

    Args:
        file: The file to read.
        file_type: The type of the file to determine the appropriate reader.

    Returns:
        str: The content of the file as a string.
    Raises:
        ValueError: If the file type is not supported.
    """

    file_readers = {
        "text/markdown": lambda f: clean_extracted_text(f.read().decode()),
        "application/pdf": read_pdf,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": read_docx,
        "text/plain": lambda f: clean_extracted_text(f.read().decode()),
    }

    if file_type not in file_readers:
        raise ValueError(f"Unsupported file type: {file_type}")

    return file_readers[file_type](file)
